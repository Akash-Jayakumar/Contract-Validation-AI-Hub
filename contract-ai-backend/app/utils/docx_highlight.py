from typing import List, Tuple
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_COLOR_INDEX

def risk_to_color(risk: int) -> WD_COLOR_INDEX:
    if risk >= 70:
        return WD_COLOR_INDEX.RED
    if risk >= 40:
        return WD_COLOR_INDEX.YELLOW
    return WD_COLOR_INDEX.BRIGHT_GREEN

def split_paragraph_by_spans(para: Paragraph, spans: List[Tuple[int, int]], color: WD_COLOR_INDEX):
    text = para.text
    if not text:
        return
    
    # Sort and merge overlapping spans
    spans = sorted([(max(0, s), max(0, e)) for s, e in spans if e > s], key=lambda x: x[0])
    
    merged = []
    for s, e in spans:
        if not merged or s > merged[-1][1]:
            merged.append([s, e])
        else:
            merged[-1][1] = max(merged[-1][1], e)
    
    # Clear the paragraph and rebuild with highlighted runs
    para.clear()
    cursor = 0
    for s, e in merged:
        if cursor < s:
            para.add_run(text[cursor:s])
        run = para.add_run(text[s:e])
        run.font.highlight_color = color
        cursor = e
    if cursor < len(text):
        para.add_run(text[cursor:])
