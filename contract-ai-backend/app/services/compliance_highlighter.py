import os
from typing import Dict, Any, List, Optional
from docx import Document
from app.utils.docx_highlight import split_paragraph_by_spans, risk_to_color

class ComplianceHighlighter:
    def __init__(self, report_title: str = "Policy Violation Report"):
        self.report_title = report_title

    def _add_header(self, doc: Document, payload: Dict[str, Any]):
        title_p = doc.add_paragraph()
        tr = title_p.add_run(self.report_title)
        tr.bold = True
        tr.font.size = 16
        
        meta = doc.add_paragraph()
        meta.add_run(f"Contract ID: {payload.get('contract_id')}\n")
        meta.add_run(f"Template: {payload.get('template_type')}  Version: {payload.get('version')}\n")
        meta.add_run(f"Overall Risk: {payload.get('overall_risk')}")

    def _sep(self, doc: Document):
        doc.add_paragraph("—" * 50)

    def _add_clause(self, doc: Document, item: Dict[str, Any], chunks: List[str]):
        title = item.get("title", item.get("policy_id", "Clause"))
        risk = int(item.get("risk", 0))
        violated = bool(item.get("violated", False))
        explanation = str(item.get("explanation", ""))
        
        p = doc.add_paragraph()
        r = p.add_run(f"{title} (risk: {risk})")
        r.bold = True
        r.font.size = 14
        
        doc.add_paragraph(explanation or "No explanation provided.")
        color = risk_to_color(risk if violated else 0)

        for ev in item.get("evidence", []):
            idx = int(ev.get("chunk_index", -1))
            if idx < 0 or idx >= len(chunks):
                continue
            text = chunks[idx] or ""
            s = int(ev.get("span", {}).get("start", 0))
            e = int(ev.get("span", {}).get("end", 0))
            s = max(0, min(s, len(text)))
            e = max(0, min(e, len(text)))
            para = doc.add_paragraph(text)
            if e > s:
                split_paragraph_by_spans(para, [(s, e)], color)

    def build_and_save(self, violation_payload: Dict[str, Any], chunks: List[str], output_path: Optional[str] = None) -> str:
        if output_path is None:
            output_path = os.path.join(os.getcwd(), "Compliance_Report.docx")
        
        doc = Document()
        self._add_header(doc, violation_payload)
        self._sep(doc)
        
        items = violation_payload.get("violations", [])
        if not items:
            doc.add_paragraph("No violations detected or insufficient evidence to assert violations.")
        else:
            for item in items:
                self._add_clause(doc, item, chunks)
                self._sep(doc)
        
        doc.save(output_path)
        return output_path
